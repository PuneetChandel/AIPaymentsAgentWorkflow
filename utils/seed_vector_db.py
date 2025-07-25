"""
Seed ChromaDB with document policies from the policies directory
"""
import os
import sys
import re
from typing import List, Dict, Any
from pathlib import Path
from dotenv import load_dotenv

# Add the project root to the Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.vector_service_optimized import VectorServiceOptimized
from utils.logging_config import init_logging, get_logger, console_print

load_dotenv()

init_logging()
logger = get_logger('utils.seed_vector_db')

class DocumentProcessor:
    """
    Process documents and extract text for ChromaDB storage
    """
    
    def __init__(self, max_chunk_size: int = 2000):
        self.max_chunk_size = max_chunk_size
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a Word document (.docx)
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Dictionary with document metadata and content
        """
        try:
            from docx import Document
            logger.info(f"Processing Word document: {file_path}")
            
            # Load document
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n".join(paragraphs)
            if tables_text:
                all_text += "\n\nTables:\n" + "\n".join(tables_text)
            
            # Generate metadata from filename and content
            metadata = self._generate_metadata(file_path, all_text)
            
            # Split into chunks if too large
            chunks = self._split_into_chunks(all_text)
            
            logger.info(f"Successfully processed {file_path}: {len(chunks)} chunks")
            
            return {
                "file_path": file_path,
                "metadata": metadata,
                "content": all_text,
                "chunks": chunks,
                "chunk_count": len(chunks)
            }
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Process all .docx files in a directory
        
        Args:
            directory_path: Path to directory containing .docx files
            
        Returns:
            List of processed document dictionaries
        """
        try:
            logger.info(f"Processing directory: {directory_path}")
            
            directory = Path(directory_path)
            if not directory.exists():
                raise FileNotFoundError(f"Directory not found: {directory_path}")
            
            # Find all .docx files
            docx_files = list(directory.glob("*.docx"))
            
            # Filter out temporary files (starting with ~$)
            docx_files = [f for f in docx_files if not f.name.startswith("~$")]
            
            logger.info(f"Found {len(docx_files)} .docx files")
            
            processed_docs = []
            for file_path in docx_files:
                try:
                    doc_data = self.process_docx(str(file_path))
                    processed_docs.append(doc_data)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
            
            logger.info(f"Successfully processed {len(processed_docs)} documents")
            return processed_docs
            
        except Exception as e:
            logger.error(f"Error processing directory {directory_path}: {e}")
            raise
    
    def _generate_metadata(self, file_path: str, content: str) -> Dict[str, Any]:
        """
        Generate metadata from filename and content
        
        Args:
            file_path: Path to the document
            content: Document content
            
        Returns:
            Dictionary with metadata
        """
        filename = Path(file_path).stem
        
        # Generate policy_id from filename
        policy_id = self._generate_policy_id(filename)
        
        # Extract title (first line or filename)
        title = self._extract_title(content, filename)
        
        # Determine category based on keywords
        category = self._determine_category(content, filename)
        
        return {
            "policy_id": policy_id,
            "title": title,
            "category": category,
            "filename": filename,
            "source": "document"
        }
    
    def _generate_policy_id(self, filename: str) -> str:
        """
        Generate a policy ID from filename
        
        Args:
            filename: The filename without extension
            
        Returns:
            Policy ID string
        """
        # Clean filename and create ID
        clean_name = re.sub(r'[^\w\s-]', '', filename)
        clean_name = re.sub(r'\s+', '-', clean_name)
        clean_name = clean_name.strip('-').upper()
        
        # Add DOC prefix
        return f"DOC-{clean_name}"
    
    def _extract_title(self, content: str, filename: str) -> str:
        """
        Extract title from content or use filename
        
        Args:
            content: Document content
            filename: Filename fallback
            
        Returns:
            Title string
        """
        lines = content.split('\n')
        
        # Look for a title in first few lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) > 5 and len(line) < 100:
                # Check if it looks like a title (not a full sentence)
                if not line.endswith('.') or line.count('.') == 0:
                    return line
        
        # Fallback to filename
        return filename.replace('_', ' ').replace('-', ' ').title()
    
    def _determine_category(self, content: str, filename: str) -> str:
        """
        Determine category based on content and filename
        
        Args:
            content: Document content
            filename: Filename
            
        Returns:
            Category string
        """
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Category keywords
        categories = {
            "Billing Dispute": ["billing", "dispute", "charge", "invoice", "payment", "refund"],
            "Service Outage": ["outage", "downtime", "service", "availability", "sla"],
            "Subscription": ["subscription", "cancel", "renewal", "plan", "upgrade"],
            "Customer Tier": ["premium", "enterprise", "vip", "tier", "priority"],
            "Fraud Prevention": ["fraud", "security", "unauthorized", "chargeback", "verification"],
            "Legal": ["legal", "terms", "agreement", "contract", "compliance"],
            "HR": ["employee", "staff", "personnel", "hiring", "benefits"],
            "General": []
        }
        
        # Check filename first
        for category, keywords in categories.items():
            if any(keyword in filename_lower for keyword in keywords):
                return category
        
        # Check content
        for category, keywords in categories.items():
            if any(keyword in content_lower for keyword in keywords):
                return category
        
        return "General"
    
    def _split_into_chunks(self, content: str) -> List[str]:
        """
        Split content into chunks if too large
        
        Args:
            content: Document content
            
        Returns:
            List of content chunks
        """
        if len(content) <= self.max_chunk_size:
            return [content]
        
        # Split by paragraphs first
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            # If single paragraph is too large, split by sentences
            if len(paragraph) > self.max_chunk_size:
                sentences = re.split(r'[.!?]+', paragraph)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    
                    if len(current_chunk) + len(sentence) + 1 > self.max_chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                            current_chunk = sentence
                        else:
                            chunks.append(sentence[:self.max_chunk_size])
                    else:
                        current_chunk += " " + sentence if current_chunk else sentence
            else:
                # Add paragraph to current chunk
                if len(current_chunk) + len(paragraph) + 2 > self.max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        current_chunk = paragraph
                    else:
                        chunks.append(paragraph)
                else:
                    current_chunk += "\n\n" + paragraph if current_chunk else paragraph
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

def reset_collections():
    """
    Reset ChromaDB collections (for testing purposes)
    """
    vector_service = VectorServiceOptimized()
    
    try:
        # Delete and recreate collections
        vector_service.client.delete_collection("dispute_resolutions")
        vector_service.client.delete_collection("company_policies")
        console_print("ChromaDB collections reset successfully", "SUCCESS")
        logger.info("ChromaDB collections reset successfully")
    except Exception as e:
        logger.warning(f"Error resetting collections (might not exist): {e}")

def load_document_policies():
    """
    Load document policies from the policies directory
    """
    policies_dir = Path("./policies")
    
    if not policies_dir.exists():
        console_print("No policies directory found. Use --create-policies-dir to create one.", "WARNING")
        return
    
    # Check if there are any .docx files
    docx_files = list(policies_dir.glob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~$")]
    
    if not docx_files:
        console_print("No .docx files found in policies directory", "WARNING")
        return
    
    console_print(f"Found {len(docx_files)} .docx files in policies directory", "INFO")
    logger.info(f"Found {len(docx_files)} .docx files in policies directory")
    
    try:
        doc_processor = DocumentProcessor()
        vector_service = VectorServiceOptimized()
        
        # Process all documents
        processed_docs = doc_processor.process_directory(str(policies_dir))
        
        if processed_docs:
            loaded_count = 0
            
            for doc_data in processed_docs:
                metadata = doc_data['metadata']
                policy_id = metadata['policy_id']
                chunks = doc_data['chunks']
                
                # Check if policy already exists
                try:
                    existing = vector_service.policies_collection.get(ids=[policy_id])
                    if existing['ids']:
                        console_print(f"Document policy {policy_id} already exists, skipping", "INFO")
                        continue
                except Exception:
                    pass
                
                # Add each chunk
                for i, chunk in enumerate(chunks):
                    chunk_id = f"{policy_id}-CHUNK-{i+1}" if len(chunks) > 1 else policy_id
                    
                    chunk_metadata = metadata.copy()
                    chunk_metadata['chunk_id'] = chunk_id
                    chunk_metadata['chunk_number'] = i + 1
                    chunk_metadata['total_chunks'] = len(chunks)
                    
                    try:
                        vector_service.policies_collection.add(
                            documents=[chunk],
                            metadatas=[chunk_metadata],
                            ids=[chunk_id]
                        )
                    except Exception as e:
                        if "existing embedding ID" not in str(e).lower():
                            logger.error(f"Error adding document policy chunk {chunk_id}: {e}")
                
                loaded_count += 1
                console_print(f"Loaded document policy: {policy_id} ({metadata['title']})", "SUCCESS")
            
            logger.info(f"Loaded {loaded_count} document policies")
            console_print(f"Loaded {loaded_count} document policies", "SUCCESS")
            
    except ImportError:
        console_print("python-docx not installed. Install with: pip install python-docx", "ERROR")
        logger.error("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"Error loading document policies: {e}")
        console_print(f"Error loading document policies: {e}", "ERROR")

def create_policies_directory():
    """
    Create the policies directory with README
    """
    policies_dir = Path("./policies")
    policies_dir.mkdir(exist_ok=True)
    
    readme_content = """# Policy Documents

This directory contains Word document policies that are loaded into ChromaDB.

## File Format
- Only .docx files are supported
- File names should be descriptive (e.g., "billing-dispute-policy.docx")
- The first line of the document is used as the title
- Category is automatically determined from filename and content

## Usage
1. Place your .docx files in this directory
2. Run: `python utils/seed_vector_db.py`
3. Files will be processed and added to ChromaDB

## Policy IDs
Policy IDs are automatically generated from filenames:
- `billing-dispute-policy.docx` → `DOC-BILLING-DISPUTE-POLICY`
- `service_outage_refunds.docx` → `DOC-SERVICE-OUTAGE-REFUNDS`

## Document Processing
- Documents are automatically chunked if they exceed 2000 characters
- Tables are extracted and included in the text
- Metadata is generated including title, category, and filename
- Duplicate documents are automatically detected and skipped

## Categories
The system automatically categorizes documents based on keywords:
- **Billing Dispute**: billing, dispute, charge, invoice, payment, refund
- **Service Outage**: outage, downtime, service, availability, sla
- **Subscription**: subscription, cancel, renewal, plan, upgrade
- **Customer Tier**: premium, enterprise, vip, tier, priority
- **Fraud Prevention**: fraud, security, unauthorized, chargeback, verification
- **Legal**: legal, terms, agreement, contract, compliance
- **HR**: employee, staff, personnel, hiring, benefits
- **General**: fallback category for uncategorized documents
"""
    
    readme_path = policies_dir / "README.md"
    with open(readme_path, 'w') as f:
        f.write(readme_content)
    
    console_print(f"Created policies directory: {policies_dir}", "SUCCESS")
    console_print(f"README created at: {readme_path}", "INFO")

def main():
    """
    Main function to seed ChromaDB with document policies
    """
    import sys
    
    # Check flags
    reset_flag = False
    create_policies_dir = False
    
    for arg in sys.argv[1:]:
        if arg == "--reset":
            reset_flag = True
        elif arg == "--create-policies-dir":
            create_policies_dir = True
    
    # Create policies directory if requested
    if create_policies_dir:
        create_policies_directory()
        return
    
    # Reset collections if requested
    if reset_flag:
        console_print("Resetting ChromaDB collections", "WARNING")
        reset_collections()
    
    console_print("Seeding ChromaDB with document policies", "SUCCESS")
    logger.info("Seeding ChromaDB with document policies")
    
    # Load document policies
    load_document_policies()
    
    console_print("ChromaDB seeding completed", "SUCCESS")
    logger.info("ChromaDB seeding completed")

if __name__ == "__main__":
    main() 